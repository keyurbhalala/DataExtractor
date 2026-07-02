"""
extractors.py — Container Data Extractor

Turns any supported shipping-paperwork file (PDF, DOCX, XLSX/XLSM, CSV, JPG, PNG)
into a normalized bundle of text blocks + images that ai_extract.py can hand to
Claude. Deliberately does NOT try to parse container numbers / weights itself —
real supplier documents vary too much in layout for hand-rolled rules to survive
contact with reality. This module's only job is faithful, format-aware
extraction of raw content.

Key behaviors:
- PDF: pdfplumber first (text-based pages). Any page with no real text layer
  (scanned page) is rasterized to a PNG image with PyMuPDF and queued for
  Claude vision instead.
- DOCX: python-docx for paragraph text + native tables, PLUS the docx is
  unzipped to check word/embeddings/ for embedded Excel objects (common when
  a supplier pastes a spreadsheet into Word instead of using a real Word
  table), PLUS word/media/ images bigger than 150px in either dimension are
  pulled as a vision fallback (tiny images are almost always logos/letterhead
  and are skipped).
- XLSX/XLSM: openpyxl, every sheet dumped as text — no assumption about which
  row is the header or which columns exist.
- CSV: decoded and passed through as raw text.
- JPG/PNG: passed straight through as vision input.
"""

from __future__ import annotations

import csv
import io
import zipfile
from dataclasses import dataclass, field
from typing import List, Tuple

import pdfplumber
import fitz  # PyMuPDF
import docx
from openpyxl import load_workbook
from PIL import Image, UnidentifiedImageError

# Minimum characters of text on a PDF page before we trust the text layer.
# Pages below this are treated as scanned/no-text and rasterized instead.
PDF_TEXT_THRESHOLD = 20

# Minimum width/height (px) for an embedded DOCX image to be treated as a
# meaningful document image (photo/scan of a table) rather than a logo.
MIN_EMBEDDED_IMAGE_DIM = 150

# DPI used when rasterizing scanned PDF pages for Claude vision.
PDF_RASTER_DPI = 200

SUPPORTED_EXTENSIONS = {
    "pdf", "docx", "xlsx", "xlsm", "csv", "jpg", "jpeg", "png",
}


@dataclass
class ImageUnit:
    data: bytes
    media_type: str  # e.g. "image/png"
    label: str        # human-readable description for debugging/logging


@dataclass
class FileExtraction:
    filename: str
    text_blocks: List[str] = field(default_factory=list)
    images: List[ImageUnit] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    @property
    def combined_text(self) -> str:
        return "\n\n".join(self.text_blocks).strip()

    @property
    def is_empty(self) -> bool:
        return not self.combined_text and not self.images


class ExtractionError(ValueError):
    """Raised when a file can't be read at all (corrupt, unsupported, etc.)."""


def extract_file(filename: str, file_bytes: bytes) -> FileExtraction:
    """Dispatch to the right extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in SUPPORTED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file type '.{ext}'. Supported: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )

    try:
        if ext == "pdf":
            return _extract_pdf(filename, file_bytes)
        if ext == "docx":
            return _extract_docx(filename, file_bytes)
        if ext in ("xlsx", "xlsm"):
            return _extract_xlsx(filename, file_bytes)
        if ext == "csv":
            return _extract_csv(filename, file_bytes)
        if ext in ("jpg", "jpeg", "png"):
            return _extract_image(filename, file_bytes)
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface as a clean per-file error
        raise ExtractionError(f"Could not read '{filename}': {exc}") from exc

    # Unreachable given the guard above, but keeps type checkers happy.
    raise ExtractionError(f"Unsupported file type '.{ext}'.")


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def _extract_pdf(filename: str, file_bytes: bytes) -> FileExtraction:
    result = FileExtraction(filename=filename)

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        n_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()

            if len(text) >= PDF_TEXT_THRESHOLD:
                block = f"--- {filename} | page {i} of {n_pages} (text layer) ---\n{text}"

                # Tables sometimes carry structure that flows badly as plain
                # text (merged cells, multi-line headers). Include them as a
                # supplementary block when pdfplumber finds any — cheap
                # insurance, Claude can ignore it if the text block above
                # already covers it.
                try:
                    tables = page.extract_tables()
                except Exception:  # noqa: BLE001
                    tables = []
                if tables:
                    table_lines = []
                    for t_idx, table in enumerate(tables, start=1):
                        table_lines.append(f"[table {t_idx}]")
                        for row in table:
                            cells = [("" if c is None else str(c).replace("\n", " ")) for c in row]
                            table_lines.append(" | ".join(cells))
                    block += (
                        f"\n\n--- {filename} | page {i} extracted tables (raw, may overlap text above) ---\n"
                        + "\n".join(table_lines)
                    )

                result.text_blocks.append(block)
            else:
                # No usable text layer -> rasterize this page for vision.
                png_bytes = _rasterize_pdf_page(file_bytes, i - 1)
                result.images.append(
                    ImageUnit(
                        data=png_bytes,
                        media_type="image/png",
                        label=f"{filename} - page {i} of {n_pages} (scanned, no text layer)",
                    )
                )

    if result.is_empty:
        result.warnings.append("No extractable text or renderable pages found.")

    return result


def _rasterize_pdf_page(file_bytes: bytes, page_index: int) -> bytes:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        page = doc[page_index]
        pix = page.get_pixmap(dpi=PDF_RASTER_DPI)
        return pix.tobytes("png")
    finally:
        doc.close()


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def _extract_docx(filename: str, file_bytes: bytes) -> FileExtraction:
    result = FileExtraction(filename=filename)
    buf = io.BytesIO(file_bytes)

    # 1) Native paragraph text + native Word tables via python-docx.
    document = docx.Document(buf)

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    if paragraphs:
        result.text_blocks.append(
            f"--- {filename} | body paragraphs ---\n" + "\n".join(paragraphs)
        )

    for t_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if any(rows):
            result.text_blocks.append(
                f"--- {filename} | native Word table {t_idx} ---\n" + "\n".join(rows)
            )

    # 2) Embedded Excel objects. A supplier will often paste a spreadsheet
    # into the Word doc as an OLE object rather than a real Word table — that
    # shows up as an .xlsx (or occasionally .xls) file under word/embeddings/
    # inside the docx zip, invisible to python-docx entirely.
    zf = zipfile.ZipFile(io.BytesIO(file_bytes))
    embedded_names = [
        n for n in zf.namelist()
        if n.startswith("word/embeddings/") and n.lower().endswith((".xlsx", ".xlsm"))
    ]
    for name in embedded_names:
        try:
            xlsx_bytes = zf.read(name)
            sheet_text = _dump_workbook_text(xlsx_bytes, source_label=f"{filename} -> embedded {name}")
            if sheet_text:
                result.text_blocks.append(sheet_text)
        except Exception as exc:  # noqa: BLE001
            result.warnings.append(f"Could not parse embedded workbook '{name}': {exc}")

    legacy_embedded = [
        n for n in zf.namelist()
        if n.startswith("word/embeddings/") and n.lower().endswith(".xls")
    ]
    for name in legacy_embedded:
        result.warnings.append(
            f"Embedded legacy .xls object '{name}' found but not parsed "
            "(only .xlsx/.xlsm embeddings are supported). Consider re-saving "
            "the source spreadsheet as .xlsx."
        )

    # 3) Embedded images >150px in either dimension, as a vision fallback
    # (e.g. a photographed weighbridge ticket pasted into the doc). Small
    # images are treated as logos/letterhead and skipped. EMF/WMF vector
    # metafiles can't be sent to a vision model and are skipped with a
    # warning.
    media_names = [n for n in zf.namelist() if n.startswith("word/media/")]
    for name in media_names:
        raw = zf.read(name)
        ext = name.rsplit(".", 1)[-1].lower()
        if ext in ("emf", "wmf"):
            result.warnings.append(
                f"Embedded image '{name}' is a {ext.upper()} vector graphic and was skipped "
                "(not viewable by the vision model)."
            )
            continue
        try:
            img = Image.open(io.BytesIO(raw))
            width, height = img.size
        except UnidentifiedImageError:
            result.warnings.append(f"Embedded file '{name}' is not a readable image and was skipped.")
            continue

        if max(width, height) < MIN_EMBEDDED_IMAGE_DIM:
            continue  # almost certainly a logo

        media_type = Image.MIME.get(img.format, "image/png")
        # Normalize to PNG bytes for consistent API payloads.
        out = io.BytesIO()
        img.convert("RGB").save(out, format="PNG")
        result.images.append(
            ImageUnit(
                data=out.getvalue(),
                media_type="image/png",
                label=f"{filename} -> embedded image {name} ({width}x{height})",
            )
        )

    if result.is_empty:
        result.warnings.append("No extractable text, embedded spreadsheet, or usable images found.")

    return result


# --------------------------------------------------------------------------
# XLSX / XLSM
# --------------------------------------------------------------------------

def _extract_xlsx(filename: str, file_bytes: bytes) -> FileExtraction:
    result = FileExtraction(filename=filename)
    text = _dump_workbook_text(file_bytes, source_label=filename)
    if text:
        result.text_blocks.append(text)
    else:
        result.warnings.append("Workbook appears to be empty.")
    return result


def _dump_workbook_text(xlsx_bytes: bytes, source_label: str) -> str:
    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True, read_only=True)
    blocks = []
    for ws in wb.worksheets:
        lines = []
        for row in ws.iter_rows(values_only=True):
            if all(c is None for c in row):
                continue
            cells = ["" if c is None else str(c) for c in row]
            lines.append(",".join(cells))
        if lines:
            blocks.append(f"--- {source_label} | sheet '{ws.title}' ---\n" + "\n".join(lines))
    return "\n\n".join(blocks)


# --------------------------------------------------------------------------
# CSV
# --------------------------------------------------------------------------

def _extract_csv(filename: str, file_bytes: bytes) -> FileExtraction:
    result = FileExtraction(filename=filename)

    text = None
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise ExtractionError(f"Could not decode '{filename}' as text (tried utf-8, latin-1).")

    # Re-serialize through csv.reader/writer just to normalize delimiters and
    # strip out any stray control characters, without assuming a fixed
    # header layout.
    try:
        rows = list(csv.reader(io.StringIO(text)))
        normalized = "\n".join(",".join(cell for cell in row) for row in rows)
    except Exception:  # noqa: BLE001 - fall back to raw text if csv module chokes
        normalized = text

    if normalized.strip():
        result.text_blocks.append(f"--- {filename} ---\n{normalized}")
    else:
        result.warnings.append("CSV file appears to be empty.")

    return result


# --------------------------------------------------------------------------
# Images (JPG/PNG)
# --------------------------------------------------------------------------

def _extract_image(filename: str, file_bytes: bytes) -> FileExtraction:
    result = FileExtraction(filename=filename)

    try:
        img = Image.open(io.BytesIO(file_bytes))
        img.verify()
    except UnidentifiedImageError as exc:
        raise ExtractionError(f"'{filename}' is not a readable image.") from exc

    ext = filename.rsplit(".", 1)[-1].lower()
    media_type = "image/jpeg" if ext in ("jpg", "jpeg") else "image/png"

    result.images.append(
        ImageUnit(data=file_bytes, media_type=media_type, label=f"{filename} (uploaded image)")
    )
    return result
